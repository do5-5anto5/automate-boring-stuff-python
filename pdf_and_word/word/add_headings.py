import docx

doc = docx.Document()

head0 = doc.add_heading('Header 0', 0)
head1 = doc.add_heading('Header 1', 1)
head2 = doc.add_heading('Header 2', 2)
head3 = doc.add_heading('Header 3', 3)
head4 = doc.add_heading('Header 4', 4)

doc.save('gen_files/headings.docx')