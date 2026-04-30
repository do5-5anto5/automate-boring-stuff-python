import docx
doc = docx.Document('gen_files/demo.docx')
doc.paragraphs[0].text
# 'Document Title'
doc.paragraphs[0].style # The exact id may be different:
# >>> _ParagraphStyle('Title') id: 3095631007984

doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text
# >>> 'A plain paragraph with some bold text and some italic'
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
# >>> ('A plain paragraph with some ', 'bold', ' and some ', 'italic')
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].undderline = True
doc.save('gen_files/restyled.docx')