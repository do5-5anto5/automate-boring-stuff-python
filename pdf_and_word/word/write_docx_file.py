import docx

doc = docx.Document()

# adds text 'Hello World' with Title style
doc.add_paragraph('Hello world.', 'Title')

paragraph2 = doc.add_paragraph('Another paragraph.')
paragraph2.add_run(' This text is being added to an existent paragraph')
doc.save('gen_files/helloworld.docx')