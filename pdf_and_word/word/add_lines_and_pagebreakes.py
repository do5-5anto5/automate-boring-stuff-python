import docx

doc = docx.Document()

doc.add_paragraph('This is in first page')

doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.add_paragraph('This is on the second page')

doc.save('gen_files/two_pages.docx')